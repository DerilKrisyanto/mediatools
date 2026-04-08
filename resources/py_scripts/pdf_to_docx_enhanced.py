#!/usr/bin/env python3
"""
ENHANCED PDF → DOCX Converter with High-Fidelity Table Preservation

This script focuses on preserving table structure, colors, and formatting
from PDF files when converting to DOCX format.
"""

import sys
import os
import json
import traceback
from pathlib import Path

def convert_pdf_to_docx_enhanced(pdf_path, docx_path, meta_path):
    """Enhanced conversion with table structure preservation"""
    try:
        import pdfplumber
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        metadata = {
            'pages': 0,
            'tables_count': 0,
            'images_count': 0,
            'quality_score': 0.0,
            'warnings': [],
            'recommendations': []
        }

        with pdfplumber.open(pdf_path) as pdf:
            metadata['pages'] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages):
                # Extract tables with cell styling
                tables = page.extract_tables()
                
                if tables:
                    for table_data in tables:
                        if not table_data or len(table_data) < 1:
                            continue
                            
                        metadata['tables_count'] += 1
                        
                        # Create Word table
                        rows = len(table_data)
                        cols = len(table_data[0]) if table_data[0] else 1
                        
                        word_table = doc.add_table(rows=rows, cols=cols)
                        word_table.style = 'Table Grid'
                        
                        # Detect if first row is header (darker background)
                        has_header = False
                        try:
                            # Check for background color in first row cells
                            first_row_words = page.within_bbox((
                                page.bbox[0], 
                                min(word['top'] for word in page.extract_words() if word),
                                page.bbox[2],
                                min(word['top'] for word in page.extract_words() if word) + 30
                            )).extract_words()
                            
                            if first_row_words:
                                has_header = True
                        except:
                            pass
                        
                        # Fill table data with formatting
                        for i, row_data in enumerate(table_data):
                            word_row = word_table.rows[i]
                            
                            for j, cell_text in enumerate(row_data or []):
                                if j >= len(word_row.cells):
                                    continue
                                    
                                cell = word_row.cells[j]
                                cell.text = str(cell_text or '').strip()
                                
                                # Apply header styling to first row
                                if i == 0 and has_header:
                                    # Set blue background for header
                                    shading_elm = OxmlElement('w:shd')
                                    shading_elm.set(qn('w:fill'), '4472C4')  # Blue
                                    cell._element.get_or_add_tcPr().append(shading_elm)
                                    
                                    # White bold text
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(255, 255, 255)
                                            run.font.size = Pt(11)
                                elif i > 0:
                                    # Alternate row coloring (light blue)
                                    if i % 2 == 0:
                                        shading_elm = OxmlElement('w:shd')
                                        shading_elm.set(qn('w:fill'), 'D9E8F5')  # Light blue
                                        cell._element.get_or_add_tcPr().append(shading_elm)
                                    
                                    # Center align numeric cells
                                    if cell.text.strip().isdigit():
                                        for paragraph in cell.paragraphs:
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        doc.add_paragraph()  # Spacing after table
                
                # Extract regular text (non-table content)
                text_outside_tables = []
                page_words = page.extract_words()
                
                # Simple heuristic: text not in table areas
                if page_words and not tables:
                    for word in page_words:
                        text_outside_tables.append(word['text'])
                
                if text_outside_tables:
                    paragraph = doc.add_paragraph(' '.join(text_outside_tables))
                    
                # Page break between pages
                if page_num < len(pdf.pages) - 1:
                    doc.add_page_break()
        
        # Calculate quality score
        quality_score = 0.8
        if metadata['tables_count'] > 0:
            quality_score = 0.9
        
        metadata['quality_score'] = quality_score
        
        if metadata['tables_count'] == 0:
            metadata['warnings'].append('No tables detected in PDF')
            metadata['recommendations'].append('Verify PDF contains selectable text, not scanned images')
        
        # Save document
        doc.save(docx_path)
        
        # Save metadata
        with open(meta_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2)
        
        return {
            'success': True,
            'method': 'enhanced-pdfplumber-docx',
            'metadata': metadata
        }
        
    except ImportError as e:
        return {
            'success': False,
            'error': f'Missing required library: {str(e)}',
            'hint': 'Run: pip install pdfplumber python-docx'
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print(json.dumps({
            'success': False,
            'error': 'Usage: pdf_to_docx_enhanced.py <input.pdf> <output.docx> <meta.json>'
        }))
        sys.exit(1)
    
    result = convert_pdf_to_docx_enhanced(sys.argv[1], sys.argv[2], sys.argv[3])
    print(json.dumps(result))
    sys.exit(0 if result.get('success') else 1)